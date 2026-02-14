# -*- coding: utf-8 -*-
"""Atualiza o roteiro com glossario de siglas"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Define cor de fundo da celula"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_glossary_table(doc, title, items, header_color="1a477a"):
    """Adiciona tabela de glossario"""
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Header
    header_cells = table.rows[0].cells
    headers = ['Sigla', 'Nome Completo', 'Significado']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
        set_cell_shading(header_cells[i], header_color)
        for run in header_cells[i].paragraphs[0].runs:
            run.font.color.rgb = None  # Will appear white on dark background

    # Rows
    for sigla, nome, significado in items:
        row_cells = table.add_row().cells
        row_cells[0].text = sigla
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = nome
        row_cells[2].text = significado

        for cell in row_cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")

def main():
    # Abrir documento existente
    doc = Document("Roteiro_Apresentacao_Qualificacao.docx")

    # Adicionar secao de glossario
    doc.add_page_break()

    title = doc.add_heading("GLOSSARIO DE SIGLAS E TERMOS TECNICOS", level=1)
    doc.add_paragraph("Referencia rapida para termos usados na apresentacao")
    doc.add_paragraph("")

    # Indices Espectrais
    indices = [
        ("NDVI", "Normalized Difference Vegetation Index", "Indice de Vegetacao por Diferenca Normalizada. Mede saude da vegetacao (-1 a +1)"),
        ("NBR", "Normalized Burn Ratio", "Razao de Queimada Normalizada. Mede severidade de queimada"),
        ("dNBR", "delta NBR", "Diferenca de NBR (antes - depois). Detecta mudanca apos fogo"),
    ]
    add_glossary_table(doc, "Indices Espectrais", indices, "2e7d32")

    # Bandas do Satelite
    bandas = [
        ("NIR", "Near Infrared", "Infravermelho Proximo (Banda 8 do Sentinel-2, 842nm)"),
        ("RED", "Red", "Vermelho visivel (Banda 4 do Sentinel-2, 665nm)"),
        ("SWIR", "Short-Wave Infrared", "Infravermelho de Onda Curta (Banda 12, 2190nm)"),
    ]
    add_glossary_table(doc, "Bandas Espectrais do Satelite", bandas, "1a477a")

    # Sensoriamento Remoto
    sensores = [
        ("FIRMS", "Fire Information for Resource Management System", "Sistema da NASA para deteccao de focos de calor em tempo quase-real"),
        ("MODIS", "Moderate Resolution Imaging Spectroradiometer", "Sensor nos satelites Terra e Aqua. Resolucao 1km"),
        ("VIIRS", "Visible Infrared Imaging Radiometer Suite", "Sensor nos satelites Suomi NPP e NOAA-20. Resolucao 375m"),
        ("MCD64A1", "MODIS Burned Area Product", "Produto mensal de area queimada. Resolucao 500m"),
        ("ERA5", "ECMWF Reanalysis 5th Generation", "Dados meteorologicos de reanalise. Resolucao ~9km"),
    ]
    add_glossary_table(doc, "Sensoriamento Remoto e Dados", sensores, "0d47a1")

    # Machine Learning
    ml = [
        ("ML", "Machine Learning", "Aprendizado de Maquina - algoritmos que aprendem com dados"),
        ("LightGBM", "Light Gradient Boosting Machine", "Algoritmo de gradient boosting otimizado para velocidade"),
        ("XGBoost", "eXtreme Gradient Boosting", "Algoritmo de gradient boosting popular em competicoes"),
        ("ROC-AUC", "Receiver Operating Characteristic - Area Under Curve", "Metrica de desempenho (0 a 1). Quanto maior, melhor"),
        ("PR-AUC", "Precision-Recall Area Under Curve", "Metrica para dados desbalanceados"),
        ("F1-Score", "F1 Score", "Media harmonica entre precisao e recall"),
        ("FP", "False Positive", "Falso Positivo - erro do tipo I"),
        ("FN", "False Negative", "Falso Negativo - erro do tipo II"),
    ]
    add_glossary_table(doc, "Machine Learning e Metricas", ml, "6a1b9a")

    # Deep Learning
    dl = [
        ("CNN", "Convolutional Neural Network", "Rede Neural Convolucional - especializada em imagens"),
        ("GPU", "Graphics Processing Unit", "Placa de video - acelera treinamento de redes neurais"),
        ("CPU", "Central Processing Unit", "Processador comum do computador"),
        ("ResNet", "Residual Network", "Arquitetura de CNN profunda com conexoes residuais"),
        ("U-Net", "U-Network", "Arquitetura para segmentacao de imagens (formato de U)"),
        ("Conv2D", "2D Convolution Layer", "Camada de convolucao bidimensional"),
        ("MaxPool", "Max Pooling", "Operacao que reduz dimensao pegando valor maximo"),
    ]
    add_glossary_table(doc, "Deep Learning (Trabalho Futuro)", dl, "e65100")

    # Termos do Projeto
    projeto = [
        ("MATOPIBA", "MA-TO-PI-BA", "Regiao agricola: Maranhao, Tocantins, Piaui, Bahia"),
        ("Weak Labeling", "Rotulacao Fraca", "Gerar rotulos automaticamente usando dados auxiliares"),
        ("Hotspot", "Ponto de Calor", "Deteccao de anomalia termica pelo satelite"),
        ("D+1", "Dia mais um", "Previsao para o dia seguinte"),
        ("Pipeline", "Fluxo de processamento", "Sequencia de etapas de processamento de dados"),
        ("Feature", "Caracteristica", "Variavel de entrada para o modelo de ML"),
    ]
    add_glossary_table(doc, "Termos do Projeto", projeto, "37474f")

    # Formulas
    doc.add_heading("Formulas dos Indices", level=2)

    formulas = [
        "NDVI = (NIR - RED) / (NIR + RED)",
        "    Valores: -1 a +1",
        "    Vegetacao saudavel: 0.6 a 0.9",
        "    Solo exposto/queimado: 0.1 a 0.3",
        "",
        "NBR = (NIR - SWIR) / (NIR + SWIR)",
        "    Valores: -1 a +1",
        "    Area queimada: valores baixos ou negativos",
        "",
        "dNBR = NBR_pre - NBR_pos",
        "    Valores positivos indicam queimada",
        "    Maior valor = maior severidade"
    ]

    for formula in formulas:
        p = doc.add_paragraph(formula)
        if formula.startswith("    "):
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].font.size = Pt(10)
            p.runs[0].italic = True
        elif "=" in formula:
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(11)

    doc.add_paragraph("")

    # Dica final
    doc.add_heading("Dica para a Apresentacao", level=2)
    tip = doc.add_paragraph()
    tip.add_run("IMPORTANTE: ").bold = True
    tip.add_run("Nao e necessario explicar todas as siglas durante a apresentacao. ")
    tip.add_run("Use este glossario para estudar e estar preparado caso a banca pergunte sobre algum termo especifico.")

    # Salvar
    doc.save("Roteiro_Apresentacao_Qualificacao.docx")
    print("Roteiro atualizado com glossario de siglas!")
    print("Arquivo: Roteiro_Apresentacao_Qualificacao.docx")

if __name__ == "__main__":
    main()
